"""
Parser import soal pilihan ganda dari file Word (.docx).

Aturan format yang didukung:
- Soal   : "1. Teks soal" / "1) Teks soal" / "Soal 1: Teks soal"
- Opsi   : "A. teks" / "a) teks" / "(B) teks" / "*C. teks" (tanda * = kunci)
- Kunci  : tanda * di depan opsi ATAU baris "Jawaban: B" / "Kunci: B"
"""
import io
import re
from typing import BinaryIO

from docx import Document

QUESTION_RE = re.compile(r"^\s*(?:soal\s*)?(\d{1,3})\s*[.)\]:\-]\s+(.+)$", re.IGNORECASE)
OPTION_RE = re.compile(r"^\s*\*?\s*\(?\s*([A-Ha-h])\s*[).\]:\-]\s+(.+)$")
ANSWER_RE = re.compile(
    r"^\s*(?:(?:kunci\s+)?jawaban|kunci|answer)\s*[:=\-]\s*\(?([A-Ha-h])\)?\s*$",
    re.IGNORECASE,
)


def _iter_block_paragraphs(document):
    """Yield paragraphs in document order, including those inside tables.
    Preserves order between normal paragraphs and tables."""
    body = document.element.body
    for child in body.iterchildren():
        # w:p = paragraph, w:tbl = table
        if child.tag.endswith('}p'):
            # find the Paragraph object matching this element
            for para in document.paragraphs:
                if para._p is child:
                    yield para
                    break
            else:
                # fallback: create temp paragraph wrapper
                from docx.text.paragraph import Paragraph
                yield Paragraph(child, document)
        elif child.tag.endswith('}tbl'):
            # find matching Table
            for table in document.tables:
                if table._tbl is child:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                yield para
                    break

def _is_list_paragraph(para):
    """Detect if paragraph is part of a Word auto-numbered/bulleted list (A., B.)"""
    try:
        pPr = para._p.pPr
        if pPr is not None and pPr.numPr is not None:
            return True
    except Exception:
        pass
    style = getattr(para.style, 'name', '') or ''
    if 'list' in style.lower():
        return True
    return False

def _normalize_line(text: str) -> str:
    if not text:
        return ''
    # hapus zero-width, NBSP, dan normalisasi spasi
    text = text.replace('\xa0', ' ').replace('\u200b', '').replace('\ufeff', '').replace('\r', ' ')
    # ganti tab dengan spasi
    text = text.replace('\t', ' ')
    return text.strip()

def parse_docx_questions(file: BinaryIO) -> dict:
    """Parse file .docx menjadi daftar soal pilihan ganda.

    Return:
        {
          "questions": [
            {
              "number": 1,
              "label": "...",
              "options": [{"label": ..., "value": ..., "order_index": 0, "is_correct": bool}],
              "errors": ["..."],   # kosong berarti soal valid
            }, ...
          ],
          "total": int,
          "valid_count": int,
        }
    """
    document = Document(file)

    questions = []
    current = None
    seen_numbers = set()
    # auto-letter untuk list tanpa huruf (Word auto-numbering)
    next_auto_letter_ord = None

    # gunakan iterator yang mencakup table
    try:
        paragraphs = list(_iter_block_paragraphs(document))
        # fallback jika iterator kosong (compat)
        if not paragraphs:
            paragraphs = document.paragraphs
    except Exception:
        paragraphs = document.paragraphs

    for para in paragraphs:
        raw = para.text or ""
        # Word soft break (Shift+Enter) menyimpan \n di dalam satu paragraph — split jadi baris terpisah
        # juga handle jika satu paragraph mengandung beberapa opsi dipisah newline
        lines = raw.splitlines() if '\n' in raw or '\r' in raw else [raw]
        for raw_line in lines:
            line = _normalize_line(raw_line)
            if not line:
                continue

            answer_match = ANSWER_RE.match(line)
            if answer_match and current is not None:
                letter = answer_match.group(1).upper()
                matched = [o for o in current["options"] if o["letter"] == letter]
                if matched:
                    for o in current["options"]:
                        o["is_correct"] = o["letter"] == letter
                else:
                    current["errors"].append(
                        f"Kunci jawaban '{letter}' tidak ada di daftar opsi"
                    )
                continue

            question_match = QUESTION_RE.match(line)
            option_match = OPTION_RE.match(line) if not question_match else None

            if question_match:
                number = int(question_match.group(1))
                current = {
                    "number": number,
                    "label": question_match.group(2).strip(),
                    "options": [],
                    "errors": [],
                }
                questions.append(current)
                if number in seen_numbers:
                    current["errors"].append(f"Nomor soal {number} duplikat")
                seen_numbers.add(number)
                next_auto_letter_ord = ord('A')
                continue

            if option_match and current is not None:
                letter = option_match.group(1).upper()
                text = option_match.group(2).strip()
                is_correct = line.lstrip().startswith("*")
                existing = next((o for o in current["options"] if o["letter"] == letter), None)
                if existing is None:
                    current["options"].append({
                        "letter": letter,
                        "label": text,
                        "value": text,
                        "order_index": len(current["options"]),
                        "is_correct": is_correct,
                    })
                else:
                    existing.update({"label": text, "value": text})
                    if is_correct:
                        existing["is_correct"] = True
                # sync auto-letter ke huruf berikutnya
                try:
                    next_auto_letter_ord = ord(letter) + 1
                except Exception:
                    pass
                continue

            # Fallback: opsi tanpa huruf karena Word auto-numbering (list A. tidak ada di text)
            if current is not None and _is_list_paragraph(para) and line and not question_match and not answer_match:
                # hanya anggap sebagai opsi jika belum banyak opsi dan teks tidak terlalu panjang untuk soal
                # dan kita sedang dalam blok opsi (sudah ada opsi sebelumnya atau baris terlihat seperti jawaban pendek)
                is_correct_fallback = line.lstrip().startswith("*")
                clean_text = line.lstrip().lstrip("*").strip()
                # hindari mengubah lanjutan soal yang panjang (>120 char) menjadi opsi
                if clean_text and len(clean_text) < 180:
                    # tentukan huruf auto
                    if next_auto_letter_ord is None:
                        next_auto_letter_ord = ord('A') + len(current["options"])
                    letter = chr(next_auto_letter_ord)
                    # jangan duplikat huruf yang sudah ada
                    if not any(o["letter"] == letter for o in current["options"]):
                        current["options"].append({
                            "letter": letter,
                            "label": clean_text,
                            "value": clean_text,
                            "order_index": len(current["options"]),
                            "is_correct": is_correct_fallback,
                        })
                        next_auto_letter_ord += 1
                        continue

            # Baris lain: anggap lanjutan teks soal
            if current is not None and not current["options"]:
                current["label"] = f"{current['label']} {line}".strip()

    result = []
    for q in questions:
        options = [{k: v for k, v in o.items() if k != "letter"} for o in q["options"]]
        errors = list(q["errors"])

        if len(options) < 2:
            errors.append("Opsi jawaban kurang dari 2")
        if not any(o["is_correct"] for o in options):
            errors.append("Kunci jawaban tidak ditemukan")

        result.append({
            "number": q["number"],
            "label": q["label"],
            "options": options,
            "errors": errors,
        })

    return {
        "questions": result,
        "total": len(result),
        "valid_count": sum(1 for q in result if not q["errors"]),
    }


def generate_template_docx() -> bytes:
    """Buat file .docx contoh template untuk di-download guru."""
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    document = Document()

    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

    for section in document.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    def set_cell_shading(cell, color_hex):
        shading_elm = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
        )
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def add_colored_heading(text, level=1, color=None):
        h = document.add_heading(text, level=level)
        if color:
            for run in h.runs:
                run.font.color.rgb = color
        return h

    BLUE = RGBColor(0x25, 0x63, 0xEB)
    DARK = RGBColor(0x1E, 0x29, 0x3B)
    GRAY = RGBColor(0x64, 0x74, 0x8B)
    GREEN = RGBColor(0x16, 0xA3, 0x4A)
    RED = RGBColor(0xDC, 0x26, 0x26)

    title = document.add_heading("Template Import Soal Pilihan Ganda", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = BLUE
        run.font.size = Pt(22)

    subtitle = document.add_paragraph("Formax — Sistem Ujian Digital")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = GRAY
        run.font.italic = True
    subtitle.paragraph_format.space_after = Pt(16)

    doc_title = document.add_paragraph()
    doc_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_line = doc_title.add_run("─" * 60)
    run_line.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)
    run_line.font.size = Pt(8)

    add_colored_heading("Petunjuk Penggunaan", level=1, color=DARK)

    petunjuk_items = [
        ("Buka form di Formax, klik tombol ", '"Import Word"', " yang ada di toolbar atas."),
        ("Download template ini lalu buka menggunakan ", "Microsoft Word", " atau ", "Google Docs", "."),
        ("Isi soal sesuai format yang sudah ditentukan di bawah ini."),
        ("Simpan file sebagai ", ".docx", " (bukan .doc lama)."),
        ("Upload file ke Formax, lalu preview dan import soal."),
    ]
    for parts in petunjuk_items:
        p = document.add_paragraph(style="List Number")
        for i, part in enumerate(parts):
            run = p.add_run(part)
            if i % 2 == 1:
                run.bold = True
                run.font.color.rgb = BLUE

    document.add_paragraph("")

    add_colored_heading("Format Penulisan Soal", level=1, color=DARK)

    format_table = document.add_table(rows=5, cols=2, style="Table Grid")
    format_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Komponen", "Format & Contoh"]
    for i, header in enumerate(headers):
        cell = format_table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(11)
        set_cell_shading(cell, "2563EB")

    rows_data = [
        ("Nomor Soal", '1. Ibu kota Indonesia adalah ...\n2) Hasil dari 15 x 15 adalah ...'),
        ("Opsi Jawaban", 'A. Bandung\nB. Jakarta\nC. Surabaya\nD. Medan'),
        ("Kunci Jawaban\n(Opsi 1)", '*B. Jakarta\n(tanda * di depan opsi yang benar)'),
        ("Kunci Jawaban\n(Opsi 2)", 'Jawaban: B\n(ditulis di baris setelah semua opsi)'),
    ]

    for row_idx, (komponen, contoh) in enumerate(rows_data, start=1):
        cell_komponen = format_table.rows[row_idx].cells[0]
        cell_komponen.text = ""
        run_k = cell_komponen.paragraphs[0].add_run(komponen)
        run_k.bold = True
        run_k.font.size = Pt(10)

        cell_contoh = format_table.rows[row_idx].cells[1]
        cell_contoh.text = ""
        run_c = cell_contoh.paragraphs[0].add_run(contoh)
        run_c.font.size = Pt(10)
        run_c.font.name = "Consolas"

        if row_idx % 2 == 0:
            set_cell_shading(cell_komponen, "F0F4FF")
            set_cell_shading(cell_contoh, "F0F4FF")

    for row in format_table.rows:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_before = Pt(4)
            cell.paragraphs[0].paragraph_format.space_after = Pt(4)

    document.add_paragraph("")

    add_colored_heading("Contoh Soal yang Benar", level=1, color=DARK)

    p_note = document.add_paragraph()
    run_note = p_note.add_run("Cara 1: Tandai kunci dengan tanda bintang (*) di depan opsi")
    run_note.bold = True
    run_note.font.color.rgb = GREEN
    run_note.font.size = Pt(11)

    samples_star = [
        ("Ibu kota Indonesia adalah ...", [
            ("A. Bandung", False), ("*B. Jakarta", True), ("C. Surabaya", False), ("D. Medan", False)
        ]),
    ]
    for qtext, opts in samples_star:
        p_q = document.add_paragraph()
        run_q = p_q.add_run(f"1. {qtext}")
        run_q.bold = True
        run_q.font.size = Pt(11)
        for opt_text, is_correct in opts:
            p_opt = document.add_paragraph()
            p_opt.paragraph_format.left_indent = Cm(1)
            run_opt = p_opt.add_run(opt_text)
            run_opt.font.size = Pt(11)
            if is_correct:
                run_opt.bold = True
                run_opt.font.color.rgb = GREEN

    document.add_paragraph("")

    p_note2 = document.add_paragraph()
    run_note2 = p_note2.add_run("Cara 2: Tulis jawaban di baris tersendiri setelah opsi")
    run_note2.bold = True
    run_note2.font.color.rgb = GREEN
    run_note2.font.size = Pt(11)

    samples_keyword = [
        ("Hasil dari 12 x 12 adalah ...", [
            ("A. 124", False), ("B. 132", False), ("C. 144", False), ("D. 154", False)
        ], "Jawaban: C"),
    ]
    for qtext, opts, answer_line in samples_keyword:
        p_q = document.add_paragraph()
        run_q = p_q.add_run(f"2. {qtext}")
        run_q.bold = True
        run_q.font.size = Pt(11)
        for opt_text, _ in opts:
            p_opt = document.add_paragraph()
            p_opt.paragraph_format.left_indent = Cm(1)
            run_opt = p_opt.add_run(opt_text)
            run_opt.font.size = Pt(11)
        p_ans = document.add_paragraph()
        p_ans.paragraph_format.left_indent = Cm(1)
        run_ans = p_ans.add_run(answer_line)
        run_ans.bold = True
        run_ans.font.color.rgb = GREEN
        run_ans.font.size = Pt(11)

    document.add_paragraph("")

    p_note3 = document.add_paragraph()
    run_note3 = p_note3.add_run("Cara 3: Gabungan (opsional)")
    run_note3.bold = True
    run_note3.font.color.rgb = GREEN
    run_note3.font.size = Pt(11)

    samples_mixed = [
        ("Planet yang dikenal sebagai planet merah adalah ...", [
            ("A. Venus", False), ("B. Bumi", False), ("C. Yupiter", False), ("*D. Mars", True)
        ], None),
    ]
    for qtext, opts, answer_line in samples_mixed:
        p_q = document.add_paragraph()
        run_q = p_q.add_run(f"3. {qtext}")
        run_q.bold = True
        run_q.font.size = Pt(11)
        for opt_text, is_correct in opts:
            p_opt = document.add_paragraph()
            p_opt.paragraph_format.left_indent = Cm(1)
            run_opt = p_opt.add_run(opt_text)
            run_opt.font.size = Pt(11)
            if is_correct:
                run_opt.bold = True
                run_opt.font.color.rgb = GREEN

    document.add_paragraph("")

    add_colored_heading("Penting!", level=2, color=RED)

    rules = [
        "Setiap soal HARUS menggunakan nomor (1. 2. 3. dst.)",
        "Opsi jawaban HARUS menggunakan huruf (A. B. C. D. dst.)",
        "Kunci jawaban WAJIB ditandai dengan salah satu cara di atas",
        "Hanya soal PILIHAN GANDA yang bisa diimport (maksimal 8 opsi)",
        "Jangan gunakan format .doc lama — simpan sebagai .docx",
    ]
    for rule in rules:
        p = document.add_paragraph(style="List Bullet")
        run = p.add_run(rule)
        run.font.size = Pt(10)

    document.add_paragraph("")

    footer_line = doc_title.add_run if False else document.add_paragraph()
    footer_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_footer = footer_line.add_run("Dibuat dengan ❤ oleh Formax — formax.com")
    run_footer.font.size = Pt(9)
    run_footer.font.color.rgb = GRAY
    run_footer.font.italic = True

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
